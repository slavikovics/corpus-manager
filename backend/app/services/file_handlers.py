import os
from typing import Dict
import logging
import PyPDF2
from docx import Document
from striprtf.striprtf import rtf_to_text
from pathlib import Path
import sharepoint2text

logger = logging.getLogger(__name__)


class FileHandler:
    SUPPORTED_ENCODINGS = ['utf-8', 'cp1251', 'latin-1', 'windows-1251', 'utf-16']

    @staticmethod
    def read_text_file(file_path: str) -> str:
        for encoding in FileHandler.SUPPORTED_ENCODINGS:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        raise ValueError(f"Could not decode file {file_path}")

    @staticmethod
    def read_pdf(file_path: str) -> str:
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                if pdf_reader.is_encrypted:
                    try:
                        pdf_reader.decrypt('')
                    except:
                        return ""

                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            logger.error(f"Error reading PDF {file_path}: {e}")
            return ""
        return text

    @staticmethod
    def read_docx(file_path: str) -> str:
        text = ""
        try:
            doc = Document(file_path)
            for paragraph in doc.paragraphs:
                if paragraph.text:
                    text += paragraph.text + "\n"

            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells if cell.text]
                    if row_text:
                        text += " | ".join(row_text) + "\n"
        except Exception as e:
            logger.error(f"Error reading DOCX {file_path}: {e}")
            return ""
        return text

    @staticmethod
    def read_doc(file_path: str) -> str:
        try:
            result = next(sharepoint2text.read_file(file_path, ignore_images=True))
            text = result.get_full_text()
            return text
        except Exception as e:
            logger.error(f"Error reading DOC: {e}")
            return ""

    @staticmethod
    def read_rtf(file_path: str) -> str:
        try:
            for encoding in FileHandler.SUPPORTED_ENCODINGS:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        rtf_content = f.read()
                        text = rtf_to_text(rtf_content)
                        if text and text.strip():
                            return text
                except UnicodeDecodeError:
                    continue
            return ""
        except Exception as e:
            logger.error(f"Error reading RTF {file_path}: {e}")
            return ""

    @classmethod
    def extract_text(cls, file_path: str, file_type: str) -> str:
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        file_type = file_type.lower().lstrip('.')

        handlers = {
            'txt': cls.read_text_file,
            'pdf': cls.read_pdf,
            'docx': cls.read_docx,
            'doc': cls.read_doc,
            'rtf': cls.read_rtf
        }

        handler = handlers.get(file_type)
        if not handler:
            raise ValueError(f"Unsupported file type: {file_type}")

        try:
            text = handler(str(file_path))
            return text if text else ""
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            return ""

    @staticmethod
    def save_upload_file_sync(upload_file, upload_dir: str, filename: str) -> str:
        os.makedirs(upload_dir, exist_ok=True)
        file_path = os.path.join(upload_dir, filename)

        counter = 1
        original_path = file_path
        while os.path.exists(file_path):
            name, ext = os.path.splitext(original_path)
            file_path = f"{name}_{counter}{ext}"
            counter += 1

        content = upload_file.read()
        with open(file_path, 'wb') as out_file:
            out_file.write(content)

        return file_path

    @staticmethod
    def read_file(file_path: str, file_type: str, metadata: Dict) -> str:
        logger.info(f"Reading file: {file_path}")

        text = FileHandler.extract_text(file_path, file_type)

        if not text or not text.strip():
            raise ValueError("File contains no extractable text")

        if not metadata.get('title'):
            lines = text.strip().split('\n')
            for line in lines:
                if line.strip():
                    metadata['title'] = line[:200].strip()
                    break

        metadata['file_size'] = os.path.getsize(file_path)
        metadata['char_count'] = len(text)
        metadata['word_count'] = len(text.split())

        return text

    @staticmethod
    async def save_upload_file(upload_file, upload_dir: str, filename: str) -> str:
        import aiofiles

        os.makedirs(upload_dir, exist_ok=True)
        file_path = os.path.join(upload_dir, filename)

        counter = 1
        original_path = file_path
        while os.path.exists(file_path):
            name, ext = os.path.splitext(original_path)
            file_path = f"{name}_{counter}{ext}"
            counter += 1

        async with aiofiles.open(file_path, 'wb') as out_file:
            content = await upload_file.read()
            await out_file.write(content)

        return file_path